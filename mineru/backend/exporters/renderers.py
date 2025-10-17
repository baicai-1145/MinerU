from __future__ import annotations

import base64
import io
import os
import re
from dataclasses import dataclass
from typing import Callable, Iterable, List, Optional, Sequence, Tuple

from bs4 import BeautifulSoup  # type: ignore[import]
from docx import Document  # type: ignore[import]
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore[import]
from docx.shared import Inches  # type: ignore[import]
from docx.oxml import parse_xml  # type: ignore[import]
from latex2mathml.converter import convert as latex_to_mathml  # type: ignore[import]
import mathml2omml  # type: ignore[import]

try:  # pylatexenc>=3.0a needed for structured sanitising
  from pylatexenc.latexnodes import nodes as latexnodes  # type: ignore[import]
  from pylatexenc.latexwalker import LatexWalker  # type: ignore[import]
  _PYLATEXENC_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency
  latexnodes = None  # type: ignore[assignment]
  LatexWalker = None  # type: ignore[assignment]
  _PYLATEXENC_AVAILABLE = False


@dataclass
class RenderAsset:
  name: str
  data: bytes
  mime: str

_ALLOWED_TEXT_COMMANDS = {
    "text",
    "mathrm",
    "mathbf",
    "mathit",
    "mathcal",
    "operatorname",
    "boldsymbol",
    "mathbb",
    "mathsf",
}
_SINGLE_CHAR_TOKENS = {"-", "+", "=", "|", "/", "(", ")", ",", ".", ":", "'"}


def _escape_html(text: str) -> str:
  return (
      text.replace("&", "&amp;")
      .replace("<", "&lt;")
      .replace(">", "&gt;")
  )


def _collapse_command_content(content: str) -> str:
  tokens = content.split()
  if not tokens:
    return content
  if all(len(token) == 1 or token in _SINGLE_CHAR_TOKENS for token in tokens):
    return "".join(tokens)
  return " ".join(tokens)


def _strip_command_spacing(body: str) -> str:
  def remove_space_after_command(match: re.Match[str]) -> str:
    return f"\\{match.group(1)}{{"

  body = re.sub(r"\\([a-zA-Z]+)\s+\{", remove_space_after_command, body)
  body = re.sub(r"([a-zA-Z0-9}])\s+([_^])", lambda m: f"{m.group(1)}{m.group(2)}", body)
  body = re.sub(r"([_^])\s+\{", r"\1{", body)
  body = re.sub(r"([_^])\s+\\", r"\1\\", body)
  body = re.sub(r"([_^])\s+", r"\1", body)
  body = re.sub(r"([a-zA-Z0-9}])\s+\\", r"\1\\", body)
  body = re.sub(r"([a-zA-Z0-9}])\s+\(", r"\1(", body)
  body = re.sub(r"\)\s+([a-zA-Z0-9\\])", r")\1", body)
  body = re.sub(r"\{\s+", "{", body)
  body = re.sub(r"\s+\}", "}", body)
  def collapse_allowed_commands(match: re.Match[str]) -> str:
    name = match.group(1)
    content = match.group(2)
    if name in _ALLOWED_TEXT_COMMANDS:
      collapsed = _collapse_command_content(content)
      return f"\\{name}{{{collapsed}}}"
    return match.group(0)

  body = re.sub(r"\\([a-zA-Z]+)\{([^{}]*)\}", collapse_allowed_commands, body)
  body = re.sub(r"\s+", " ", body)
  body = body.strip()
  body = re.sub(r"\\s+\{", "{", body)
  return body


def _normalise_math_tokens(body: str) -> str:
  body = _strip_command_spacing(body)
  body = re.sub(r"\\left\s*\(", r"\\left(", body)
  body = re.sub(r"\\right\s*\)", r"\\right)", body)
  body = re.sub(r"\\sum\s*_{", r"\\sum_{", body)
  body = re.sub(r"\\int\s*_{", r"\\int_{", body)
  body = re.sub(r"\\mathcal\s*\{", r"\\mathcal{", body)
  body = re.sub(r"\\mathrm\s*\{", r"\\mathrm{", body)
  body = re.sub(r"\\boldsymbol\s*\{", r"\\boldsymbol{", body)
  return body


def _extract_math(text: str) -> Tuple[str, bool]:
  body = text.strip()
  is_display = False
  if body.startswith("$$") and body.endswith("$$"):
    is_display = True
    body = body[2:-2].strip()
  elif body.startswith("\\[") and body.endswith("\\]"):
    is_display = True
    body = body[2:-2].strip()
  elif body.startswith("\\(") and body.endswith("\\)"):
    body = body[2:-2].strip()

  body = body.replace("\n", " ")
  body = _normalise_math_tokens(body)
  body = _sanitize_with_pylatexenc(body)
  return body, is_display


_OMML_NAMESPACE = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_LATEX_TAG_PATTERN = re.compile(r"\\tag\s*\{([^{}]*)\}")
_LATEX_REPLACEMENTS = (
    (r"\pmb", r"\boldsymbol"),
    (r"\mathbf ", r"\mathbf"),
    (r"\boldsymbol ", r"\boldsymbol"),
)
_LATEX_SYMBOL_MAP = {
    r"\\rightarrow": "→",
    r"\\leftarrow": "←",
    r"\\Rightarrow": "⇒",
    r"\\Leftarrow": "⇐",
    r"\\leftrightarrow": "↔",
    r"\\mapsto": "↦",
    r"\\to": "→",
    r"\\leq": "≤",
    r"\\geq": "≥",
    r"\\times": "×",
    r"\\pm": "±",
    r"\\cdot": "·",
    r"\\neq": "≠",
}

_PYLATEX_MACROS = {name for name in [
    "mathcal",
    "mathrm",
    "mathbf",
    "boldsymbol",
    "mathit",
    "operatorname",
    "text",
    "mathsf",
    "mathbb",
]}


if _PYLATEXENC_AVAILABLE:
  def _latexnodes_to_string(nodelist) -> str:
    parts: List[str] = []
    for node in nodelist:
      if isinstance(node, latexnodes.LatexCharsNode):
        parts.append(node.chars)
      elif isinstance(node, latexnodes.LatexGroupNode):
        parts.append("{")
        parts.append(_latexnodes_to_string(node.nodelist))
        parts.append("}")
      elif isinstance(node, latexnodes.LatexMacroNode):
        parts.append(f"\\{node.macroname}")
        args = node.nodeargd.argnlist if node.nodeargd else []
        for arg in args:
          if arg is None:
            continue
          content = _latexnodes_to_string(arg.nodelist)
          if node.macroname in _PYLATEX_MACROS:
            content = _collapse_command_content(content)
          parts.append("{")
          parts.append(content)
          parts.append("}")
        if node.macro_post_space:
          parts.append(node.macro_post_space.strip())
      elif isinstance(node, latexnodes.LatexMathNode):
        left, right = node.delimiters
        parts.append(left or "")
        parts.append(_latexnodes_to_string(node.nodelist))
        parts.append(right or "")
      else:
        try:
          parts.append(node.latex_verbatim())
        except Exception:
          continue
    return "".join(parts)
else:  # pragma: no cover - pylatexenc unavailable
  def _latexnodes_to_string(nodelist) -> str:
    return ""


def _sanitize_with_pylatexenc(expr: str) -> str:
  if not _PYLATEXENC_AVAILABLE:
    return expr
  try:
    walker = LatexWalker(expr)  # type: ignore[call-arg]
    nodelist, _, _ = walker.get_latex_nodes()
    cleaned = _latexnodes_to_string(nodelist)
    cleaned = cleaned.replace("\\ ", "\\")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned or expr
  except Exception:  # pragma: no cover - safety net
    return expr


def _split_latex_tag(expr: str) -> tuple[str, Optional[str]]:
  tag_value: Optional[str] = None

  def _capture(match: re.Match[str]) -> str:
    nonlocal tag_value
    tag_value = match.group(1).strip()
    return ""

  cleaned = _LATEX_TAG_PATTERN.sub(_capture, expr)
  return cleaned.strip(), tag_value


def _prepare_latex(expr: str) -> str:
  prepared = expr.strip()
  for src, dst in _LATEX_REPLACEMENTS:
    prepared = prepared.replace(src, dst)
  return prepared


def _latex_to_plaintext(expr: str) -> str:
  text = expr
  for pattern, replacement in _LATEX_SYMBOL_MAP.items():
    text = text.replace(pattern, replacement)
  text = re.sub(r"\\(mathrm|text|operatorname|mathsf|mathbf|boldsymbol)\s*\{([^{}]*)\}",
                lambda m: _collapse_command_content(m.group(2)),
                text)
  text = text.replace(r"\,", " ")
  text = text.replace(r"\ ", " ")
  return text


def _latex_to_omml_xml(latex_expr: str, *, inline: bool) -> str:
  mathml = latex_to_mathml(latex_expr)
  omml_body = mathml2omml.convert(mathml).strip()
  if not omml_body.startswith("<m:oMath"):
    raise ValueError("unexpected OMML output")
  if "xmlns:m" not in omml_body:
    omml_body = omml_body.replace("<m:oMath", f'<m:oMath xmlns:m="{_OMML_NAMESPACE}"', 1)
  if inline:
    return omml_body
  return f'<m:oMathPara xmlns:m="{_OMML_NAMESPACE}">{omml_body}</m:oMathPara>'


def _append_docx_equation(document: Document, latex_expr: str) -> bool:
  cleaned, tag_value = _split_latex_tag(latex_expr)
  if not cleaned:
    return False
  prepared = _prepare_latex(cleaned)
  try:
    omml_xml = _latex_to_omml_xml(prepared, inline=False)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph._p.append(parse_xml(omml_xml))
    if tag_value:
      paragraph.add_run(f" ({tag_value})")
    return True
  except Exception:  # pragma: no cover - graceful degradation
    return False


def _append_inline_math(paragraph, latex_expr: str) -> bool:
  cleaned, tag_value = _split_latex_tag(latex_expr)
  if not cleaned:
    return False
  prepared = _prepare_latex(cleaned)
  try:
    omml_xml = _latex_to_omml_xml(prepared, inline=True)
    paragraph._p.append(parse_xml(omml_xml))
    if tag_value:
      paragraph.add_run(f" ({tag_value})")
    return True
  except Exception:
    fallback_expr, fallback_tag = _split_latex_tag(latex_expr)
    fallback_expr = _latex_to_plaintext(fallback_expr)
    paragraph.add_run(fallback_expr if not fallback_tag else f"{fallback_expr} ({fallback_tag})")
    return False


def _split_math_segments(text: str) -> Sequence[Tuple[str, str]]:
  pattern = re.compile(r"(\$\$.*?\$\$|\$.*?\$)", re.DOTALL)
  parts: List[Tuple[str, str]] = []
  last = 0
  for match in pattern.finditer(text):
    if match.start() > last:
      parts.append(("TEXT", text[last:match.start()]))
    parts.append(("MATH", match.group(0)))
    last = match.end()
  if last < len(text):
    parts.append(("TEXT", text[last:]))
  return parts


def _escape_latex_segment(segment: str) -> str:
  replacements = {
      "\\": r"\textbackslash{}",
      "&": r"\&",
      "%": r"\%",
      "$": r"\$",
      "#": r"\#",
      "_": r"\_",
      "{": r"\{",
      "}": r"\}",
      "~": r"\textasciitilde{}",
      "^": r"\^{}",
  }
  for target, repl in replacements.items():
    segment = segment.replace(target, repl)
  return segment


def _normalize_inline_html(text: str) -> str:
  soup = BeautifulSoup(text, "html.parser")
  for sup in soup.find_all("sup"):
    sup.insert_before("^(")
    sup.insert_after(")")
    sup.decompose()
  for sub in soup.find_all("sub"):
    sub.insert_before("_(")
    sub.insert_after(")")
    sub.decompose()
  cleaned = soup.get_text()
  # collapse whitespace but preserve single spaces
  cleaned = re.sub(r"\s+", " ", cleaned)
  return cleaned.strip()


def latex_escape_text(text: str) -> str:
  parts = _split_math_segments(text)
  escaped: list[str] = []
  for kind, segment in parts:
    if kind == "MATH":
      escaped.append(segment)
    else:
      escaped.append(_escape_latex_segment(segment))
  return "".join(escaped)


def content_list_to_html(
    content_list: Sequence[dict],
    image_loader: Callable[[str], Optional[RenderAsset]],
) -> str:
  body: List[str] = []
  for item in content_list:
    item_type = item.get("type")
    if item_type == "text":
      level = item.get("text_level")
      raw = item.get("text", "")
      if level and 1 <= level <= 4:
        body.append(f"<h{level}>{raw}</h{level}>")
      else:
        html_text = raw.replace("\n", "<br/>")
        body.append(f"<p>{html_text}</p>")
    elif item_type == "equation":
      content = item.get("text", "")
      math_body, is_display = _extract_math(content)
      if not math_body:
        continue
      body.append('<div class="equation">')
      if is_display or "\\tag" in math_body:
        wrapped = f"$${math_body}$$"
      else:
        wrapped = f"${math_body}$"
      body.append(_escape_html(wrapped))
      body.append("</div>")
    elif item_type == "list":
      list_items = item.get("list_items") or []
      if not list_items:
        continue
      ordered = all(re.match(r"^\s*\d+[\.\)]", entry) for entry in list_items)
      tag = "ol" if ordered else "ul"
      body.append(f"<{tag}>")
      for entry in list_items:
        body.append(f"<li>{entry}</li>")
      body.append(f"</{tag}>")
    elif item_type == "image":
      path = item.get("img_path")
      if not path:
        continue
      asset = image_loader(path)
      if not asset:
        continue
      data_url = f"data:{asset.mime};base64,{base64.b64encode(asset.data).decode()}"
      body.append('<figure class="image">')
      body.append(f'<img src="{data_url}" alt="figure"/>')
      captions = item.get("image_caption") or []
      if captions:
        body.append(f"<figcaption>{'<br/>'.join(captions)}</figcaption>")
      body.append("</figure>")
    elif item_type == "table":
      table_html = item.get("table_body")
      if table_html:
        body.append(table_html)
      elif item.get("img_path"):
        asset = image_loader(item["img_path"])
        if asset:
          data_url = f"data:{asset.mime};base64,{base64.b64encode(asset.data).decode()}"
          body.append(f'<img src="{data_url}" alt="table"/>')
      captions = item.get("table_caption") or []
      if captions:
        body.append(f"<p class=\"table-caption\">{'<br/>'.join(captions)}</p>")
    elif item_type == "code":
      code_body = item.get("code_body") or ""
      lang = item.get("guess_lang") or ""
      captions = item.get("code_caption") or []
      if captions:
        body.append(f"<p class=\"code-caption\">{'<br/>'.join(captions)}</p>")
      body.append(f"<pre><code class=\"language-{lang}\">{_escape_html(code_body)}</code></pre>")
    else:
      # ignore headers/footers etc
      continue

  html_parts = [
      "<!DOCTYPE html>",
      "<html lang=\"zh-CN\">",
      "<head>",
      "  <meta charset=\"utf-8\" />",
      "  <title>MinerU Export</title>",
      "  <style>",
      "    body { font-family: 'Helvetica Neue', Arial, 'Microsoft YaHei', sans-serif; padding: 2rem; line-height: 1.6; }",
      "    h1, h2, h3, h4 { font-weight: 600; margin-top: 2em; }",
      "    .equation { text-align: center; margin: 1.5em 0; font-family: 'CMU Serif', 'Times New Roman', serif; }",
      "    figure.image { text-align: center; margin: 1.5em auto; }",
      "    figure.image img { max-width: 90%; height: auto; }",
      "    figcaption { font-size: 0.9em; color: #555; margin-top: 0.5em; }",
      "    .table-caption { font-size: 0.9em; color: #555; margin: 0.5em 0 1.5em 0; }",
      "    pre { background: #1f1f1f; color: #f8f8f2; padding: 1rem; border-radius: 0.5rem; overflow-x: auto; }",
      "    table { border-collapse: collapse; width: 100%; margin: 1.5em 0; }",
      "    table, th, td { border: 1px solid #ddd; }",
      "    th, td { padding: 0.75em; text-align: left; }",
      "  </style>",
      "  <script>",
      "    window.MathJax = {",
      "      tex: {",
      r"        inlineMath: [['$', '$'], ['\(', '\)']],",
      r"        displayMath: [['$$', '$$'], ['\[', '\]']],",
      "      },",
      "      svg: { fontCache: 'global' }",
      "    };",
      "  </script>",
      "  <script async src=\"https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js\"></script>",
      "</head>",
      "<body>",
      *body,
      "</body>",
      "</html>",
  ]
  return "\n".join(html_parts)


def _html_to_plain_text(text: str) -> str:
  soup = BeautifulSoup(text, "html.parser")
  for br in soup.find_all("br"):
    br.replace_with("\n")
  return soup.get_text()


def content_list_to_docx(
    content_list: Sequence[dict],
    image_loader: Callable[[str], Optional[RenderAsset]],
) -> bytes:
  document = Document()
  for item in content_list:
    item_type = item.get("type")
    if item_type == "text":
      raw_text = item.get("text", "")
      segments = _split_math_segments(raw_text)
      level = item.get("text_level")
      if len(segments) == 1 and segments[0][0] == "TEXT":
        text = _html_to_plain_text(segments[0][1])
        if not text.strip():
          continue
        if level and 1 <= level <= 4:
          document.add_heading(text, level=min(level, 4))
        else:
          paragraph = document.add_paragraph(text)
          paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        continue

      if level and 1 <= level <= 4:
        paragraph = document.add_heading("", level=min(level, 4))
      else:
        paragraph = document.add_paragraph()

      appended = False
      for kind, segment in segments:
        if kind == "TEXT":
          plain = _html_to_plain_text(segment)
          if plain:
            paragraph.add_run(plain)
            appended = True
        else:
          math_body, _ = _extract_math(segment)
          if not math_body:
            continue
          _append_inline_math(paragraph, math_body)
          appended = True

      if not appended:
        paragraph._p.getparent().remove(paragraph._p)
        continue

      if not (level and 1 <= level <= 4):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    elif item_type == "equation":
      text = item.get("text", "")
      math_body, _ = _extract_math(text)
      if not math_body:
        continue
      if not _append_docx_equation(document, math_body):
        fallback_expr, fallback_tag = _split_latex_tag(math_body)
        fallback_expr = _latex_to_plaintext(fallback_expr)
        fallback_text = fallback_expr if not fallback_tag else f"{fallback_expr} ({fallback_tag})"
        fallback = document.add_paragraph(fallback_text or math_body)
        fallback.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif item_type == "list":
      list_items = item.get("list_items") or []
      ordered = all(re.match(r"^\s*\d+[\.\)]", entry) for entry in list_items)
      style = "List Number" if ordered else "List Bullet"
      for entry in list_items:
        document.add_paragraph(_html_to_plain_text(entry), style=style)
    elif item_type == "image":
      path = item.get("img_path")
      if not path:
        continue
      asset = image_loader(path)
      if not asset:
        continue
      image_stream = io.BytesIO(asset.data)
      document.add_picture(image_stream, width=Inches(6))
      captions = item.get("image_caption") or []
      if captions:
        cap_para = document.add_paragraph(" ".join(captions))
        cap_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif item_type == "table":
      table_html = item.get("table_body")
      if table_html:
        soup = BeautifulSoup(table_html, "html.parser")
        rows = soup.find_all("tr")
        if rows:
          doc_table = document.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))
          doc_table.style = "Light List Accent 1"
          for r_idx, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for c_idx, cell in enumerate(cells):
              doc_table.rows[r_idx].cells[c_idx].text = _html_to_plain_text(cell.decode_contents()).strip()
      elif item.get("img_path"):
        asset = image_loader(item["img_path"])
        if asset:
          document.add_picture(io.BytesIO(asset.data), width=Inches(6))
      captions = item.get("table_caption") or []
      if captions:
        paragraph = document.add_paragraph(" ".join(captions))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif item_type == "code":
      code_body = item.get("code_body") or ""
      captions = item.get("code_caption") or []
      if captions:
        document.add_paragraph(" ".join(captions))
      paragraph = document.add_paragraph()
      run = paragraph.add_run(code_body)
      font = run.font
      font.name = "Courier New"
    else:
      continue

  buffer = io.BytesIO()
  document.save(buffer)
  return buffer.getvalue()


def content_list_to_latex(
    content_list: Sequence[dict],
) -> Tuple[str, List[str]]:
  lines: List[str] = []
  used_images: List[str] = []
  for item in content_list:
    item_type = item.get("type")
    if item_type == "text":
      text = item.get("text", "")
      level = item.get("text_level")
      normalized = _normalize_inline_html(text)
      escaped = latex_escape_text(normalized)
      if level == 1:
        lines.append(f"\\section{{{escaped}}}")
      elif level == 2:
        lines.append(f"\\subsection{{{escaped}}}")
      elif level == 3:
        lines.append(f"\\subsubsection{{{escaped}}}")
      elif level == 4:
        lines.append(f"\\paragraph{{{escaped}}}")
      else:
        lines.append(f"{escaped}\n")
    elif item_type == "equation":
      text = item.get("text", "")
      math_body, _ = _extract_math(text)
      if "\\tag" in math_body:
        lines.append("\\begin{equation}")
        lines.append(f"  {math_body}")
        lines.append("\\end{equation}")
      else:
        lines.append("\\[")
        lines.append(f"  {math_body}")
        lines.append("\\]")
    elif item_type == "list":
      list_items = item.get("list_items") or []
      ordered = all(re.match(r"^\\s*\\d+[\\.\\)]", entry) for entry in list_items)
      env = "enumerate" if ordered else "itemize"
      lines.append(f"\\begin{{{env}}}")
      for entry in list_items:
        normalized = _normalize_inline_html(entry)
        escaped = latex_escape_text(normalized)
        lines.append(f"  \\item {escaped}")
      lines.append(f"\\end{{{env}}}")
    elif item_type == "image":
      path = item.get("img_path")
      if not path:
        continue
      used_images.append(path)
      caption_parts = item.get("image_caption") or []
      lines.append("\\begin{figure}[htbp]")
      lines.append("  \\centering")
      lines.append(f"  \\includegraphics[width=0.8\\linewidth]{{{path}}}")
      if caption_parts:
        caption_text = latex_escape_text(" ".join(caption_parts))
        lines.append(f"  \\caption{{{caption_text}}}")
      lines.append("\\end{figure}")
    elif item_type == "table":
      table_html = item.get("table_body")
      if table_html:
        soup = BeautifulSoup(table_html, "html.parser")
        rows = soup.find_all("tr")
        if rows:
          col_count = len(rows[0].find_all(["td", "th"]))
          lines.append("\\begin{tabular}{" + "|".join(["l"] * col_count) + "}")
          lines.append("\\hline")
          for row in rows:
            cells = row.find_all(["td", "th"])
            cell_texts = [latex_escape_text(_normalize_inline_html(cell.decode_contents())) for cell in cells]
            lines.append(" & ".join(cell_texts) + " \\\\ \\hline")
          lines.append("\\end{tabular}")
      elif item.get("img_path"):
        img_path = item["img_path"]
        used_images.append(img_path)
        lines.append("\\begin{figure}[htbp]")
        lines.append("  \\centering")
        lines.append(f"  \\includegraphics[width=0.8\\linewidth]{{{img_path}}}")
        caption_parts = item.get("table_caption") or []
        if caption_parts:
          caption_text = latex_escape_text(" ".join(caption_parts))
          lines.append(f"  \\caption{{{caption_text}}}")
        lines.append("\\end{figure}")
    elif item_type == "code":
      code_body = item.get("code_body") or ""
      escaped = latex_escape_text(code_body)
      lines.append("\\begin{verbatim}")
      lines.append(escaped)
      lines.append("\\end{verbatim}")

  document = [
      "\\documentclass{article}",
      "\\usepackage[UTF8]{ctex}",
      "\\usepackage{amsmath, amssymb}",
      "\\usepackage{graphicx}",
      "\\usepackage{float}",
      "\\usepackage{microtype}",
      "\\usepackage{hyperref}",
      "\\setlength{\\emergencystretch}{3em}",
      "\\sloppy",
      "\\begin{document}",
  ]
  document.extend(lines)
  document.append("\\end{document}")
  return "\n".join(document), used_images
